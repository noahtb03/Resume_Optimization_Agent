"""DOCX export: TailoredResume -> clean, ATS-friendly Word document.

Cross-platform (pure python-docx, no Word/Node/LibreOffice needed for .docx).
Deliberately ATS-safe: single column, no tables/text-boxes/graphics, standard
headings, real bullet list formatting, Arial. Metrics already live inside the
bullet text (validator guarantees the display string is present), so nothing
special is needed to render them.

One-page fitting lives in fit.py and is applied before this renderer by the
export service.
"""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from ..models.output import TailoredResume

_DATE_MONTHS = {
    1: "Jan", 2: "Feb", 3: "Mar", 4: "Apr", 5: "May", 6: "Jun",
    7: "Jul", 8: "Aug", 9: "Sep", 10: "Oct", 11: "Nov", 12: "Dec",
}


def _fmt_date(dm) -> str:
    if dm is None:
        return "Present"
    if dm.month and dm.month in _DATE_MONTHS:
        return f"{_DATE_MONTHS[dm.month]} {dm.year}"
    return str(dm.year)


def _date_range(start, end) -> str:
    s = _fmt_date(start) if start else ""
    e = _fmt_date(end)  # end=None -> "Present"
    if not s:
        return e
    return f"{s} \u2013 {e}"


def _setup_styles(doc: Document, base_pt: float = 9.0) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(base_pt)
    # tighten default paragraph spacing for a compact resume
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0


def _add_hyperlink(paragraph, url: str, text: str, size_pt: float) -> None:
    """Add a clickable hyperlink run with short display text (keeps the contact
    line on one line instead of printing full URLs)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))  # half-points
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _link_label(url: str) -> str:
    """Short display label for a known link, else a cleaned host."""
    u = url.lower()
    if "github.com" in u:
        return "GitHub"
    if "linkedin.com" in u:
        return "LinkedIn"
    if "noahbennettdev" in u or "portfolio" in u:
        return "Portfolio"
    # generic: strip scheme + www + trailing slash
    label = url.split("://")[-1].lstrip("www.").rstrip("/")
    return label


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # a thin bottom border via paragraph border (no tables)
    _bottom_border(p)


def _bottom_border(paragraph) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.size = Pt(9)


def render_docx(resume: TailoredResume, out_path: str) -> str:
    doc = Document()
    _setup_styles(doc, base_pt=9.0)

    section = doc.sections[0]
    # US Letter + 0.7in margins for a compact one-page-friendly layout
    from docx.shared import Inches
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(0.7))

    c = resume.contact

    # --- header: name + contact line ---
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    nrun = name_p.add_run(c.full_name)
    nrun.bold = True
    nrun.font.size = Pt(16)

    text_bits = [b for b in (c.phone, c.email, c.location) if b]
    if text_bits or c.links:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(4)
        first = True
        for b in text_bits:
            if not first:
                cp.add_run("  |  ").font.size = Pt(8.5)
            cp.add_run(b).font.size = Pt(8.5)
            first = False
        for link in c.links:
            if not first:
                cp.add_run("  |  ").font.size = Pt(8.5)
            _add_hyperlink(cp, link, _link_label(link), 8.5)
            first = False

    # --- sections in declared order ---
    for sec in resume.section_order:
        if sec == "summary" and resume.summary and resume.summary.text.strip():
            _add_heading(doc, "Summary")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.add_run(resume.summary.text).font.size = Pt(9)

        elif sec == "skills" and resume.skills:
            _add_heading(doc, "Skills")
            # group by category, comma-joined (no tables)
            by_cat: dict[str, list[str]] = {}
            for sk in resume.skills:
                by_cat.setdefault(sk.category or "Other", []).append(sk.name)
            for cat, names in by_cat.items():
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                r = p.add_run(f"{cat}: ")
                r.bold = True
                r.font.size = Pt(9)
                p.add_run(", ".join(names)).font.size = Pt(9)

        elif sec == "experience" and resume.experiences:
            _add_heading(doc, "Experience")
            for exp in resume.experiences:
                _entry_header(doc, f"{exp.title}, {exp.employer}",
                              exp.location, _date_range(exp.start, exp.end))
                for b in exp.bullets:
                    _add_bullet(doc, b.text)

        elif sec == "projects" and resume.projects:
            _add_heading(doc, "Projects")
            for pr in resume.projects:
                _entry_header(doc, pr.name, None, "")
                for b in pr.bullets:
                    _add_bullet(doc, b.text)

        elif sec == "education" and resume.education:
            _add_heading(doc, "Education")
            for ed in resume.education:
                line = ed.degree + (f", {ed.field}" if ed.field else "")
                _entry_header(doc, line, ed.institution,
                              _date_range(ed.start, ed.end) if ed.end else "")
                if ed.details:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    p.add_run("  |  ".join(ed.details)).font.size = Pt(8.5)

        elif sec == "certifications" and resume.certifications:
            _add_heading(doc, "Certifications")
            for cert in resume.certifications:
                issuer = f" \u2014 {cert.issuer}" if cert.issuer else ""
                p = doc.add_paragraph()
                p.add_run(f"{cert.name}{issuer}").font.size = Pt(9)

    doc.save(out_path)
    return out_path


def _entry_header(doc: Document, left: str, location, right: str) -> None:
    """A bold title line with a right-aligned date via a tab stop (no tables)."""
    from docx.shared import Inches
    from docx.enum.text import WD_TAB_ALIGNMENT
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    # right tab stop at the content edge (7.1in usable for 0.7in margins)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.1), WD_TAB_ALIGNMENT.RIGHT)
    lr = p.add_run(left)
    lr.bold = True
    lr.font.size = Pt(9.5)
    if right:
        p.add_run("\t")
        dr = p.add_run(right)
        dr.font.size = Pt(8.5)
    if location:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(1)
        loc = lp.add_run(location)
        loc.italic = True
        loc.font.size = Pt(8.5)
